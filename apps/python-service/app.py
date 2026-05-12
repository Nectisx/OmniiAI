"""
OmniAI — Microservice Python Flask
Fonctions:
  - Extraction texte PDF (PyMuPDF)
  - Extraction texte DOCX (python-docx)
  - Speech-to-Text (whisper / SpeechRecognition)
"""
import os
import io
import logging
from functools import wraps
from flask import Flask, request, jsonify
from flask_cors import CORS

# Configuration logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app, origins=[os.getenv('BACKEND_URL', 'http://localhost:3001')])

MAX_TEXT_CHARS = 50_000  # RG16: limite 50k caractères


# ── Middleware sécurité ─────────────────────────────────────

def require_service_secret(f):
    """Vérifie le secret interne pour les appels service-à-service."""
    @wraps(f)
    def decorated(*args, **kwargs):
        secret = request.headers.get('X-Service-Secret', '')
        expected = os.getenv('PYTHON_SERVICE_SECRET', '')
        if expected and secret != expected:
            logger.warning(f"Tentative accès non autorisé depuis {request.remote_addr}")
            return jsonify({'success': False, 'error': 'Non autorisé'}), 401
        return f(*args, **kwargs)
    return decorated


# ── Health check ────────────────────────────────────────────

@app.route('/health', methods=['GET'])
def health():
    return jsonify({
        'success': True,
        'service': 'OmniAI Python Service',
        'version': '1.0.0',
        'status': 'healthy'
    })


# ── Extraction texte PDF ─────────────────────────────────────

def extract_pdf_text(file_bytes: bytes) -> str:
    """Extrait le texte d'un PDF avec PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype='pdf')
        text_parts = []
        for page_num, page in enumerate(doc):
            text = page.get_text('text')
            if text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
        doc.close()
        return '\n\n'.join(text_parts)
    except ImportError:
        # Fallback: PyPDF2 si PyMuPDF non disponible
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(f"--- Page {i + 1} ---\n{text}")
            return '\n\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Erreur extraction PDF: {e}")
            raise


def extract_docx_text(file_bytes: bytes) -> str:
    """Extrait le texte d'un fichier Word DOCX."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Inclure les tableaux
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(' | '.join(cells))

        return '\n\n'.join(paragraphs)
    except Exception as e:
        logger.error(f"Erreur extraction DOCX: {e}")
        raise


# ── Route extraction de texte ────────────────────────────────

@app.route('/extract', methods=['POST'])
@require_service_secret
def extract_text():
    """
    POST /extract
    Reçoit un fichier et retourne son texte extrait.
    Supporte: PDF, DOCX, TXT
    Limite: 50 000 caractères (RG16)
    """
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'Aucun fichier reçu'}), 400

    file = request.files['file']
    if not file.filename:
        return jsonify({'success': False, 'error': 'Nom de fichier manquant'}), 400

    filename = file.filename.lower()
    file_bytes = file.read()

    logger.info(f"Extraction: {file.filename} ({len(file_bytes)} bytes)")

    try:
        if filename.endswith('.pdf'):
            text = extract_pdf_text(file_bytes)
        elif filename.endswith('.docx'):
            text = extract_docx_text(file_bytes)
        elif filename.endswith('.txt'):
            text = file_bytes.decode('utf-8', errors='replace')
        else:
            return jsonify({
                'success': False,
                'error': f'Format non supporté: {filename.split(".")[-1]}'
            }), 400

        # RG16: Limiter à 50k caractères
        truncated = len(text) > MAX_TEXT_CHARS
        if truncated:
            text = text[:MAX_TEXT_CHARS]
            logger.info(f"Texte tronqué à {MAX_TEXT_CHARS} chars")

        return jsonify({
            'success': True,
            'text': text,
            'chars': len(text),
            'truncated': truncated,
            'filename': file.filename
        })

    except Exception as e:
        logger.error(f"Erreur extraction {filename}: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ── Speech-to-Text ───────────────────────────────────────────

@app.route('/transcribe', methods=['POST'])
@require_service_secret
def transcribe_audio():
    """
    POST /transcribe
    Reçoit un fichier audio et retourne la transcription.
    Utilise OpenAI Whisper (local) ou SpeechRecognition.
    """
    if 'audio' not in request.files:
        return jsonify({'success': False, 'error': 'Aucun fichier audio reçu'}), 400

    audio_file = request.files['audio']
    audio_bytes = audio_file.read()
    language = request.form.get('language', 'fr')

    logger.info(f"Transcription: {audio_file.filename} ({len(audio_bytes)} bytes), lang={language}")

    # Tentative Whisper (plus précis)
    try:
        import whisper
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix='.webm', delete=False) as tmp:
            tmp.write(audio_bytes)
            tmp_path = tmp.name

        try:
            model = whisper.load_model('base')
            result = model.transcribe(tmp_path, language=language)
            return jsonify({
                'success': True,
                'text': result['text'].strip(),
                'language': result.get('language', language),
                'engine': 'whisper'
            })
        finally:
            os.unlink(tmp_path)

    except ImportError:
        pass  # Whisper non installé, fallback

    # Fallback: SpeechRecognition + Google Speech API
    try:
        import speech_recognition as sr
        import tempfile
        import os

        r = sr.Recognizer()

        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp:
            tmp.write(audio_bytes)
            tmp_path = tmp.name

        try:
            with sr.AudioFile(tmp_path) as source:
                audio = r.record(source)

            lang_code = 'fr-FR' if language == 'fr' else 'en-US'
            text = r.recognize_google(audio, language=lang_code)

            return jsonify({
                'success': True,
                'text': text,
                'language': language,
                'engine': 'google_speech'
            })
        finally:
            os.unlink(tmp_path)

    except Exception as e:
        logger.error(f"Erreur transcription: {e}")
        return jsonify({'success': False, 'error': f'Transcription échouée: {str(e)}'}), 500


# ── Info service ─────────────────────────────────────────────

@app.route('/info', methods=['GET'])
def info():
    """Retourne les capacités du service selon les bibliothèques disponibles."""
    capabilities = {
        'pdf': False,
        'docx': False,
        'whisper': False,
        'speech_recognition': False
    }

    try:
        import fitz
        capabilities['pdf'] = True
    except ImportError:
        try:
            import PyPDF2
            capabilities['pdf'] = True
        except ImportError:
            pass

    try:
        from docx import Document
        capabilities['docx'] = True
    except ImportError:
        pass

    try:
        import whisper
        capabilities['whisper'] = True
    except ImportError:
        pass

    try:
        import speech_recognition
        capabilities['speech_recognition'] = True
    except ImportError:
        pass

    return jsonify({'success': True, 'capabilities': capabilities})


# ── Démarrage ────────────────────────────────────────────────

if __name__ == '__main__':
    port = int(os.getenv('PYTHON_PORT', 5000))
    debug = os.getenv('FLASK_ENV', 'production') == 'development'
    logger.info(f"🐍 OmniAI Python Service démarré sur le port {port}")
    app.run(host='0.0.0.0', port=port, debug=debug)
